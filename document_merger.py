import os
import tempfile
from typing import List
import requests
from copy import deepcopy
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import streamlit as st
from doc_converter import DocConverter

class DocumentMerger:
    """Handle merging of Word documents containing clauses"""
    
    def __init__(self, template_path: str = "clauses/Exemple contrat V2 clausier km.docx", enable_summary: bool = False):
        self.output_dir = tempfile.mkdtemp()
        self.template_path = template_path
        self.doc_converter = DocConverter()
        self.enable_summary = enable_summary
    
    def merge_documents(self, file_paths: List[str], clause_names: List[str]) -> str:
        """
        Merge multiple Word documents into one using template as base
        
        Args:
            file_paths: List of paths to Word documents to merge
            clause_names: List of clause names for headers
            
        Returns:
            Path to the merged document
        """
        if not file_paths:
            raise ValueError("Aucun document à fusionner")
        
        # Use template as base document
        if os.path.exists(self.template_path):
            final_doc = Document(self.template_path)
        else:
            st.warning(f"Template non trouvé: {self.template_path}, utilisation d'un document vide")
            final_doc = Document()
            
        # Find insertion point (after existing content)
        insertion_point = len(final_doc.paragraphs)
        
        for i, (file_path, clause_name) in enumerate(zip(file_paths, clause_names)):
            try:
                # Add section header with custom styling and insert after it
                header_para = self._add_section_header(final_doc, clause_name.upper())
                
                # Optional spacing paragraph after header
                spacer_para = final_doc.add_paragraph()
                anchor_el = spacer_para._element
                
                # Read source document
                source_doc = Document(file_path)
                
                # Insert source content right after the header area
                anchor_el = self._insert_document_body_after(final_doc, anchor_el, source_doc)
                
                # Add page break between clauses (except for the last one)
                if i < len(file_paths) - 1:
                    final_doc.add_page_break()
                    
            except Exception as e:
                st.warning(f"Erreur lors de la fusion de {clause_name}: {str(e)}")
                continue
        
        # Save merged document
        output_path = os.path.join(self.output_dir, 'document_final.docx')
        final_doc.save(output_path)
        
        status_text.text("Fusion terminée!")
        # Try to generate and embed a brief summary
        if self.enable_summary:
            try:
                summary = self.summarize_document(output_path)
                if summary:
                    st.info("📝 Résumé automatique généré et inséré en tête de document.")
                    # Save alongside document as .txt for convenience
                    with open(os.path.join(self.output_dir, 'document_final_summary.txt'), 'w', encoding='utf-8') as f:
                        f.write(summary)
                    # Embed near template marker ("Synthèse") if found, otherwise at top
                    doc = Document(output_path)
                    if not self._insert_summary_after_marker(doc, summary):
                        self._insert_summary_at_top(doc, summary)
                    doc.save(output_path)
            except Exception as _e:
                pass
        return output_path
    
    def merge_documents_by_sections(self, clauses_by_section: dict, sections_order: list) -> str:
        """
        Merge documents organized by contract sections
        
        Args:
            clauses_by_section: Dictionary of section_key -> list of clause objects
            sections_order: List of section objects in order
            
        Returns:
            Path to the merged document
        """
        # Use template as base document
        if os.path.exists(self.template_path):
            final_doc = Document(self.template_path)
        else:
            st.warning(f"Template non trouvé: {self.template_path}, utilisation d'un document vide")
            final_doc = Document()
        
        total_sections = len([s for s in sections_order if clauses_by_section.get(s['key'], [])])
        current_section = 0
        displayed_index = 0  # Dynamic numbering counter for displayed sections only
        
        for section in sections_order:
            section_clauses = clauses_by_section.get(section['key'], [])
            
            if not section_clauses:
                continue
                
            current_section += 1
            displayed_index += 1
            
            # Add section title and create an anchor right after
            header_para = self._add_section_header(final_doc, f"{displayed_index}. {section['name'].upper()}")
            spacer_para = final_doc.add_paragraph()
            anchor_el = spacer_para._element
            
            # Process each clause in this section
            for clause in section_clauses:
                try:
                    # Read source document with error handling for different formats
                    source_doc = self._safe_load_document(clause['file_path'])
                    
                    # Insert content after the current anchor, then advance anchor
                    anchor_el = self._insert_document_body_after(final_doc, anchor_el, source_doc)
                    
                    # Add spacing paragraph and move anchor to it
                    spacer_para = final_doc.add_paragraph()
                    anchor_el = spacer_para._element
                    
                except Exception as e:
                    st.warning(f"Erreur lors de la fusion de {clause['name']}: {str(e)}")
                    continue
            
            # Add spacing between sections (except for the last one)
            if current_section < total_sections:
                final_doc.add_paragraph()  # First line break
                final_doc.add_paragraph()  # Second line break
            
        # Save merged document
        output_path = os.path.join(self.output_dir, 'document_final.docx')
        final_doc.save(output_path)
        # Try to generate and embed a brief summary
        if self.enable_summary:
            try:
                summary = self.summarize_document(output_path)
                if summary:
                    st.info("📝 Résumé automatique généré et inséré en tête de document.")
                    with open(os.path.join(self.output_dir, 'document_final_summary.txt'), 'w', encoding='utf-8') as f:
                        f.write(summary)
                    doc = Document(output_path)
                    if not self._insert_summary_after_marker(doc, summary):
                        self._insert_summary_at_top(doc, summary)
                    doc.save(output_path)
            except Exception:
                pass
        return output_path
    
    def add_table_of_contents(self, doc: Document, clause_names: List[str]):
        """Add a simple table of contents"""
        toc_heading = doc.add_heading('Table des Matières', level=1)
        
        for i, clause_name in enumerate(clause_names, 1):
            toc_para = doc.add_paragraph(f"{i}. {clause_name}")
            toc_para.style = 'List Number'
        
        doc.add_page_break()
    
    def _add_section_header(self, doc: Document, section_name: str):
        """Add a section header with Montserrat ExtraBold uppercase styling"""
        para = doc.add_paragraph()
        run = para.add_run(section_name)
        
        # Set font to Montserrat ExtraBold
        run.font.name = 'Montserrat ExtraBold'
        run.font.size = Pt(11)  # Changed to 11pt
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x3D, 0xA5)  # #003DA5
        
        # Set paragraph alignment
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add spacing after for line break before first clause
        para.space_after = Pt(12)
        para.space_before = Pt(24)
    
    def _add_clause_content(self, doc: Document, content: str):
        """Deprecated: kept for compatibility; prefer _append_document_body."""
        para = doc.add_paragraph(content)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.space_after = Pt(6)
        para.space_before = Pt(0)

    def _append_document_body(self, target: Document, source: Document) -> None:
        """Append block elements from source into target, preserving formatting, enforcing template font, and avoiding section copy.

        Also imports numbering definitions from source and remaps paragraph list numIds to keep bullets/numbered lists.
        """
        from docx.oxml.ns import qn
        target_body = target._body._element
        appended_blocks = []
        numid_map = self._import_numbering_and_build_map(target, source)

        for block in list(source._body._element):
            tag = block.tag.rsplit('}', 1)[-1]
            # Skip section properties to avoid columns/margins overriding template
            if tag == 'sectPr':
                continue

            clone = deepcopy(block)
            # Also strip any per-paragraph section properties
            if tag == 'p':
                pPr = clone.find(qn('w:pPr'))
                if pPr is not None:
                    sectPr = pPr.find(qn('w:sectPr'))
                    if sectPr is not None:
                        pPr.remove(sectPr)
                    # Remap numbering id if present
                    numPr = pPr.find(qn('w:numPr'))
                    if numPr is not None:
                        numId = numPr.find(qn('w:numId'))
                        if numId is not None:
                            old_val = numId.get(qn('w:val'))
                            if old_val in numid_map:
                                numId.set(qn('w:val'), str(numid_map[old_val]))

            target_body.append(clone)
            appended_blocks.append(clone)

        # Apply Montserrat blue style and clean fields on appended content only
        self._apply_montserrat_and_clean_fields(target, appended_blocks)

    def _insert_document_body_after(self, target: Document, anchor_element, source: Document):
        """Insert source blocks right after a given anchor paragraph element and return the last inserted element to be used as new anchor."""
        from docx.oxml.ns import qn
        body = target._body._element
        last_inserted = anchor_element
        inserted_blocks = []
        numid_map = self._import_numbering_and_build_map(target, source)

        for block in list(source._body._element):
            tag = block.tag.rsplit('}', 1)[-1]
            if tag == 'sectPr':
                continue
            clone = deepcopy(block)
            if tag == 'p':
                pPr = clone.find(qn('w:pPr'))
                if pPr is not None:
                    sectPr = pPr.find(qn('w:sectPr'))
                    if sectPr is not None:
                        pPr.remove(sectPr)
                    # Remap numbering id if present
                    numPr = pPr.find(qn('w:numPr'))
                    if numPr is not None:
                        numId = numPr.find(qn('w:numId'))
                        if numId is not None:
                            old_val = numId.get(qn('w:val'))
                            if old_val in numid_map:
                                numId.set(qn('w:val'), str(numid_map[old_val]))
            # Insert after anchor
            body.insert(body.index(last_inserted) + 1, clone)
            last_inserted = clone
            inserted_blocks.append(clone)

        # Apply styling/cleanup to inserted blocks
        self._apply_montserrat_and_clean_fields(target, inserted_blocks)
        return last_inserted

    def _apply_montserrat_and_clean_fields(self, doc: Document, appended_blocks: list) -> None:
        """Force Montserrat blue font on appended content while preserving bold/italics and remove page-number fields."""
        from docx.oxml.ns import qn
        appended_set = set(appended_blocks)

        # Paragraphs
        for para in doc.paragraphs:
            if para._element in appended_set:
                # Fallback normalization for lists that use symbol fonts or bullet glyphs
                self._fallback_convert_symbol_bullets(para)
                for run in para.runs:
                    # Remove field instruction text like PAGE/NUMPAGES
                    for child in list(run._element):
                        tag = child.tag.rsplit('}', 1)[-1]
                        if tag == 'instrText':
                            text_val = (child.text or '').upper()
                            if any(key in text_val for key in [' PAGE ', 'NUMPAGES', 'PAGEREF', 'TOC']):
                                child.text = ''
                        elif tag == 'fldChar':
                            # Clear run if it's part of a field
                            run.text = ''
                    # Enforce template font (keep bold/italic as-is)
                    run.font.name = 'Montserrat Medium'
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0x00, 0x3D, 0xA5)

        # Tables
        from docx.table import _Cell, Table
        # Build map of appended table elements
        appended_tables = {blk for blk in appended_set if blk.tag.rsplit('}', 1)[-1] == 'tbl'}
        if appended_tables:
            for table in doc.tables:
                if table._element in appended_tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                self._fallback_convert_symbol_bullets(p)
                                for run in p.runs:
                                    for child in list(run._element):
                                        tag = child.tag.rsplit('}', 1)[-1]
                                        if tag == 'instrText':
                                            text_val = (child.text or '').upper()
                                            if any(key in text_val for key in [' PAGE ', 'NUMPAGES', 'PAGEREF', 'TOC']):
                                                child.text = ''
                                        elif tag == 'fldChar':
                                            run.text = ''
                                    run.font.name = 'Montserrat Medium'
                                    run.font.size = Pt(11)
                                    run.font.color.rgb = RGBColor(0x00, 0x3D, 0xA5)


    def _import_numbering_and_build_map(self, target: Document, source: Document) -> dict:
        """Import numbering definitions from source into target and return a map of source numId -> new target numId.

        This allows pasted paragraphs to keep their list/bullet numbering without visual glitches.
        """
        try:
            t_num_part = target.part.numbering_part
            s_num_part = source.part.numbering_part
        except Exception:
            return {}

        t_num_el = t_num_part.element
        s_num_el = s_num_part.element

        # Collect existing IDs in target
        from docx.oxml.ns import qn
        def existing_ids(root, tag_name):
            return {int(el.get(qn('w:%sId' % ('abstractNum' if tag_name == 'abstractNum' else 'num')))) for el in root.findall(qn('w:%s' % tag_name))}

        t_abs_ids = existing_ids(t_num_el, 'abstractNum') if t_num_el is not None else set()
        t_num_ids = existing_ids(t_num_el, 'num') if t_num_el is not None else set()

        next_abs_id = (max(t_abs_ids) + 1) if t_abs_ids else 1
        next_num_id = (max(t_num_ids) + 1) if t_num_ids else 1

        numid_map: dict[str, int] = {}

        # Map of source abstractNum id to new target abstractNum id
        abs_map: dict[str, int] = {}

        # Copy abstractNum definitions
        for abs_el in s_num_el.findall(qn('w:abstractNum')) if s_num_el is not None else []:
            old_abs_id = abs_el.get(qn('w:abstractNumId'))
            new_abs_id = next_abs_id
            next_abs_id += 1
            abs_clone = deepcopy(abs_el)
            abs_clone.set(qn('w:abstractNumId'), str(new_abs_id))
            t_num_el.append(abs_clone)
            abs_map[old_abs_id] = new_abs_id

        # Copy num definitions and track numId mapping
        for num_el in s_num_el.findall(qn('w:num')) if s_num_el is not None else []:
            old_num_id = num_el.get(qn('w:numId'))
            # Find its abstractNumId
            abs_id_el = num_el.find(qn('w:abstractNumId'))
            old_abs_id = abs_id_el.get(qn('w:val')) if abs_id_el is not None else None
            new_abs_id = abs_map.get(old_abs_id)
            new_num_id = next_num_id
            next_num_id += 1
            num_clone = deepcopy(num_el)
            num_clone.set(qn('w:numId'), str(new_num_id))
            if abs_id_el is not None:
                # Update abstractNumId reference
                abs_in_clone = num_clone.find(qn('w:abstractNumId'))
                if abs_in_clone is not None and new_abs_id is not None:
                    abs_in_clone.set(qn('w:val'), str(new_abs_id))
            t_num_el.append(num_clone)
            numid_map[old_num_id] = new_num_id

        return numid_map

    def _fallback_convert_symbol_bullets(self, para) -> None:
        """Detect paragraphs that visually look like bullets via symbol glyphs (•, ·, ◦, ✓, ✔, etc.)
        or Wingdings/Symbol fonts and convert to a safe template bullet/number list.

        Only triggers if no numbering is present on the paragraph.
        """
        from docx.oxml.ns import qn
        import re
        # Skip if paragraph already has numbering
        p_el = para._element
        pPr = p_el.find(qn('w:pPr'))
        if pPr is not None and pPr.find(qn('w:numPr')) is not None:
            return

        # Build raw text and inspect first run for symbol fonts/glyphs
        raw = para.text or ''
        stripped = raw.strip()

        bullet_glyphs = ('•', '◦', '·', '▪', '–', '-', '—', '‣', '∙', '○', '■', '□', '✓', '✔')
        starts_with_bullet = stripped.startswith(bullet_glyphs)

        symbol_fonts = {'Symbol', 'Wingdings', 'Wingdings 2', 'Wingdings 3', 'Webdings'}
        first_run = para.runs[0] if para.runs else None
        run_has_symbol_font = bool(first_run and first_run.font and first_run.font.name in symbol_fonts)

        # Numeric artifact like "1/3 " at the start
        ratio_prefix = re.match(r"^\s*\d+\/\d+\s+", raw)

        if starts_with_bullet or run_has_symbol_font or ratio_prefix:
            # Clean leading markers
            text = stripped
            # remove ratio like 1/3
            text = re.sub(r"^\d+\/\d+\s+", "", text)
            # remove common bullet/line markers
            text = text.lstrip('•').lstrip('◦').lstrip('·').lstrip('▪').lstrip('–').lstrip('-').lstrip('—').lstrip('‣').lstrip('∙').lstrip('○').lstrip('■').lstrip('□').lstrip('✓').lstrip('✔').lstrip('\t ').lstrip()

            # Replace paragraph text with cleaned text
            try:
                # Clear all runs and set text once to avoid duplicated markers
                for r in list(para.runs):
                    r.clear()
                para.add_run(text)
            except Exception:
                para.text = text

            # Apply bullet style by default; if it looks like an ordered item (starts with digits + . or )) then List Number
            if re.match(r"^\d+[\.)]\s+", stripped):
                try:
                    para.style = 'List Number'
                except Exception:
                    pass
            else:
                try:
                    para.style = 'List Bullet'
                except Exception:
                    pass
    
    def _safe_load_document(self, file_path: str) -> Document:
        """Safely load a Word document with format detection and automatic conversion"""
        try:
            # Try loading as standard Word document
            return Document(file_path)
        
        except Exception as e:
            # Check if it's a legacy .doc file that needs conversion
            if file_path.endswith('.doc') and self.doc_converter.is_legacy_doc_file(file_path):
                # Generate placeholder for unsupported .doc files
                converted_path = self.doc_converter.convert_doc_to_docx(file_path)
                
                if converted_path and os.path.exists(converted_path):
                    try:
                        # Load the converted document - this now contains a placeholder
                        converted_doc = Document(converted_path)
                        return converted_doc
                    except Exception as conv_error:
                        st.error(f"❌ Erreur lors du traitement: {str(conv_error)}")
                        raise ValueError(f"Impossible de traiter le fichier: {str(conv_error)}")
                else:
                    raise ValueError(f"Impossible de traiter le fichier: {os.path.basename(file_path)}")
            
            # If it's not a legacy .doc file or conversion failed, provide helpful error message
            import mimetypes
            mime_type, _ = mimetypes.guess_type(file_path)
            
            error_msg = f"Impossible de lire le fichier '{os.path.basename(file_path)}'"
            
            if "not a Word file" in str(e):
                error_msg += ". Le fichier semble être dans un format non supporté."
            elif "content type" in str(e):
                error_msg += ". Le fichier a une extension incorrecte ou est corrompu."
                if "themeManager" in str(e):
                    error_msg += " Il semble être un fichier Office mal nommé."
            else:
                error_msg += f". Erreur: {str(e)}"
            
            raise ValueError(error_msg)

    def summarize_document(self, docx_path: str, max_chars: int = 16000) -> str:
        """Create a short summary in French of the generated contract using OpenAI API key from cleAPI.txt."""
        try:
            text = self._read_docx_text(docx_path, max_chars=max_chars)
        except Exception as e:
            raise e

        api_key = self._read_api_key()
        if not api_key:
            return ""

        payload = {
            "model": "gpt-4o-mini",
            "temperature": 0.3,
            "max_tokens": 500,
            "messages": [
                {"role": "system", "content": "Tu es un assistant juridique. Produis des synthèses structurées, claires et factuelles en français. Utilise uniquement du texte brut, sans formatting markdown."},
                {"role": "user", "content": f"Analyse le contrat suivant et produis une synthèse structurée en texte brut avec :\n\nPOINTS CLÉS\n• [Liste des éléments essentiels du contrat]\n• [Un point par ligne avec des puces simples]\n\nCONFLITS\n• [0 à 3 points sur d'éventuelles contradictions entre clauses]\n• [Ou indiquer \"Aucun conflit de clause détecté\" si tu n'en trouves pas]\n\nUtilise uniquement des caractères simples (•) pour les listes, pas de markdown.\n\nTexte:\n{text}"}
            ]
        }

        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }

        resp = requests.post("https://api.openai.com/v1/chat/completions", json=payload, headers=headers, timeout=60)
        resp.raise_for_status()
        data = resp.json()
        content = data.get("choices", [{}])[0].get("message", {}).get("content", "").strip()
        return content

    def _read_docx_text(self, path: str, max_chars: int = 16000) -> str:
        doc = Document(path)
        parts = []
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                parts.append(p.text.strip())
            if sum(len(x) for x in parts) >= max_chars:
                break
        return "\n".join(parts)[:max_chars]

    def _read_api_key(self) -> str:
        """Read API key from cleAPI.txt in project root or Streamlit secrets."""
        # First try Streamlit secrets (works in cloud deployments)
        try:
            import streamlit as st
            if hasattr(st, 'secrets') and 'OPENAI_API_KEY' in st.secrets:
                key = st.secrets['OPENAI_API_KEY']
                if key and key.strip():
                    return key.strip()
        except Exception:
            pass
        
        # Then try environment variable
        try:
            import os
            key = os.getenv('OPENAI_API_KEY')
            if key and key.strip():
                return key.strip()
        except Exception:
            pass
        
        # Finally try local file
        candidates = [
            os.path.join(os.getcwd(), 'cleAPI.txt'),
            os.path.abspath(os.path.join(os.path.dirname(__file__), 'cleAPI.txt')),
            os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'cleAPI.txt')),
        ]
        for p in candidates:
            try:
                if os.path.exists(p):
                    with open(p, 'r', encoding='utf-8') as f:
                        key = f.read().strip()
                        if key:
                            return key
            except Exception:
                continue
        return ""

    def _insert_summary_at_top(self, doc: Document, summary_text: str) -> None:
        """Insert a 'Synthèse' section at the very beginning of the document with styled heading and body."""
        from docx.oxml.ns import qn
        body = doc._body._element
        # Create heading paragraph
        heading = Document().add_paragraph()  # temp paragraph
        h_run = heading.add_run("Synthèse")
        h_run.font.name = 'Montserrat ExtraBold'
        h_run.font.size = Pt(11)
        h_run.font.bold = True
        h_run.font.color.rgb = RGBColor(0x00, 0x3D, 0xA5)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading.space_after = Pt(6)
        heading_el = deepcopy(heading._element)
        # Insert heading at index 0
        body.insert(0, heading_el)
        # Create body paragraph(s)
        for line in summary_text.split('\n'):
            p = Document().add_paragraph()
            r = p.add_run(line)
            
            # Apply bold formatting to section titles
            line_upper = line.upper().strip()
            if line_upper in ['POINTS CLÉS', 'POINTS CLES', 'CONFLITS']:
                r.font.name = 'Montserrat ExtraBold'
                r.font.bold = True
                p.space_before = Pt(6)
            else:
                r.font.name = 'Montserrat Medium'
            
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0x00, 0x3D, 0xA5)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.space_after = Pt(3)
            body.insert(1, deepcopy(p._element))
        # Add a blank paragraph after summary
        spacer = Document().add_paragraph()
        body.insert(1, deepcopy(spacer._element))

    def _insert_summary_after_marker(self, doc: Document, summary_text: str) -> bool:
        """Insert summary right after the template marker line containing 'Synthèse' (accent aware) if present.

        Returns True if inserted at marker, else False.
        """
        import unicodedata
        def normalize(s: str) -> str:
            return ''.join(c for c in unicodedata.normalize('NFD', s.lower()) if unicodedata.category(c) != 'Mn')

        target_tokens = {"synthese", "synthese:"}
        paragraphs = list(doc.paragraphs)
        insert_idx = None
        for i, p in enumerate(paragraphs):
            text = (p.text or '').strip()
            if not text:
                continue
            norm = normalize(text)
            if any(tok in norm for tok in target_tokens):
                insert_idx = i + 1
                # If next line is an empty placeholder line, skip it to insert after
                if insert_idx < len(paragraphs) and not (paragraphs[insert_idx].text or '').strip():
                    insert_idx += 1
                break

        if insert_idx is None:
            return False

        # Build heading and body like at top, but without adding another heading if the marker already is the heading
        body = doc._body._element
        # Create body paragraph(s)
        from copy import deepcopy as _dc
        # We will insert starting from insert_idx; compute XML element reference at that position
        # If insert_idx >= current length, append at end
        if insert_idx >= len(paragraphs):
            # Append a small spacer first
            spacer = Document().add_paragraph()
            body.append(_dc(spacer._element))
            for line in summary_text.split('\n'):
                p = Document().add_paragraph()
                r = p.add_run(line)
                
                # Apply bold formatting to section titles
                line_upper = line.upper().strip()
                if line_upper in ['POINTS CLÉS', 'POINTS CLES', 'CONFLITS']:
                    r.font.name = 'Montserrat ExtraBold'
                    r.font.bold = True
                    p.space_before = Pt(6)
                else:
                    r.font.name = 'Montserrat Medium'
                
                r.font.size = Pt(11)
                r.font.color.rgb = RGBColor(0x00, 0x3D, 0xA5)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.space_after = Pt(3)
                body.append(_dc(p._element))
            return True

        anchor_el = paragraphs[insert_idx]._element
        # Insert a spacer then summary paragraphs before the anchor_el
        spacer = Document().add_paragraph()
        body.insert(body.index(anchor_el), _dc(spacer._element))

        for line in summary_text.split('\n'):
            p = Document().add_paragraph()
            r = p.add_run(line)
            
            # Apply bold formatting to section titles
            line_upper = line.upper().strip()
            if line_upper in ['POINTS CLÉS', 'POINTS CLES', 'CONFLITS']:
                r.font.name = 'Montserrat ExtraBold'
                r.font.bold = True
                p.space_before = Pt(6)
            else:
                r.font.name = 'Montserrat Medium'
            
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0x00, 0x3D, 0xA5)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.space_after = Pt(3)
            body.insert(body.index(anchor_el), _dc(p._element))

        return True
    
    def cleanup(self):
        """Clean up temporary files"""
        try:
            import shutil
            shutil.rmtree(self.output_dir, ignore_errors=True)
            # Also cleanup converter temporary files
            self.doc_converter.cleanup()
        except Exception:
            pass