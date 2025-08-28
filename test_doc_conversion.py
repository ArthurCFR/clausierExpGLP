#!/usr/bin/env python3

from doc_converter import DocConverter
from local_client import LocalClauseClient
import os
import types

def mock_streamlit():
    """Mock streamlit functions for testing"""
    def mock_info(msg):
        print(f"Info: {msg}")
    
    def mock_success(msg):
        print(f"Success: {msg}")
    
    def mock_warning(msg):
        print(f"Warning: {msg}")
    
    def mock_error(msg):
        print(f"Error: {msg}")
    
    # Mock streamlit functions
    import streamlit as st
    st.info = mock_info
    st.success = mock_success
    st.warning = mock_warning
    st.error = mock_error

def test_doc_conversion():
    print("🔄 Test de conversion des fichiers .doc legacy")
    print("=" * 60)
    
    mock_streamlit()
    
    # Test the converter directly
    converter = DocConverter()
    
    # Look for .doc files in the clauses directory
    doc_files = []
    for root, dirs, files in os.walk("clauses"):
        for file in files:
            if file.endswith('.doc'):
                doc_files.append(os.path.join(root, file))
    
    print(f"📄 {len(doc_files)} fichier(s) .doc trouvé(s):")
    for doc_file in doc_files:
        print(f"   - {doc_file}")
    
    if not doc_files:
        print("❌ Aucun fichier .doc trouvé pour le test")
        return
    
    # Test each .doc file
    for doc_file in doc_files:
        print(f"\n🔍 Test du fichier: {doc_file}")
        
        # Check if it's a legacy file
        is_legacy = converter.is_legacy_doc_file(doc_file)
        print(f"   Legacy format: {'✅ Oui' if is_legacy else '❌ Non'}")
        
        if is_legacy:
            print(f"   🔄 Tentative de conversion...")
            try:
                converted_path = converter.convert_doc_to_docx(doc_file)
                if converted_path and os.path.exists(converted_path):
                    print(f"   ✅ Conversion réussie: {converted_path}")
                    size = os.path.getsize(converted_path)
                    print(f"   📊 Taille du fichier converti: {size} octets")
                    
                    # Try to read the converted file
                    from docx import Document
                    try:
                        doc = Document(converted_path)
                        para_count = len(doc.paragraphs)
                        text_paragraphs = [p for p in doc.paragraphs if p.text.strip()]
                        print(f"   📝 Paragraphes: {para_count} (dont {len(text_paragraphs)} avec texte)")
                        
                        # Show first few lines of content
                        print(f"   📖 Aperçu du contenu:")
                        for i, para in enumerate(text_paragraphs[:3]):
                            preview = para.text[:80] + "..." if len(para.text) > 80 else para.text
                            print(f"      {i+1}. {preview}")
                            
                    except Exception as e:
                        print(f"   ❌ Erreur lors de la lecture du fichier converti: {e}")
                        
                else:
                    print(f"   ❌ Échec de la conversion")
                    
            except Exception as e:
                print(f"   ❌ Erreur durant la conversion: {e}")
    
    print(f"\n" + "=" * 60)
    
    # Test with LocalClauseClient
    print("🧪 Test avec LocalClauseClient")
    
    client = LocalClauseClient()
    clause_files = client.get_clause_files()
    
    legacy_clauses = [f for f in clause_files if f['file_path'].endswith('.doc')]
    print(f"📄 {len(legacy_clauses)} clause(s) .doc détectée(s) par le client:")
    
    for clause in legacy_clauses:
        print(f"   - {clause['name']} ({clause['section_name']})")
    
    # Cleanup
    converter.cleanup()
    client.cleanup()
    
    print("✅ Test terminé")

if __name__ == "__main__":
    test_doc_conversion()