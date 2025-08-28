#!/usr/bin/env python3

from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def analyze_template():
    template_path = "clauses/Exemple contrat V2 clausier km.docx"
    
    if not os.path.exists(template_path):
        print(f"❌ Template non trouvé: {template_path}")
        return
    
    print("🔍 Analyse du document template")
    print("=" * 50)
    
    try:
        doc = Document(template_path)
        
        print(f"📄 Nombre de paragraphes: {len(doc.paragraphs)}")
        print(f"📊 Nombre de tableaux: {len(doc.tables)}")
        
        print("\n📝 Premiers paragraphes:")
        for i, para in enumerate(doc.paragraphs[:10]):
            if para.text.strip():
                print(f"  {i+1}. {para.text[:100]}{'...' if len(para.text) > 100 else ''}")
                
                # Analyze style
                if para.runs:
                    run = para.runs[0]
                    font = run.font
                    print(f"     Style: {font.name}, Taille: {font.size}, Gras: {font.bold}")
        
        print("\n📋 Styles utilisés:")
        styles_found = set()
        for para in doc.paragraphs:
            if para.style.name:
                styles_found.add(para.style.name)
        
        for style in sorted(styles_found):
            print(f"  - {style}")
        
        print("\n🎨 Polices utilisées:")
        fonts_found = set()
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.name:
                    fonts_found.add(run.font.name)
        
        for font in sorted(fonts_found):
            print(f"  - {font}")
            
        print("\n" + "=" * 50)
        print("✅ Analyse terminée")
        
    except Exception as e:
        print(f"❌ Erreur lors de l'analyse: {e}")

if __name__ == "__main__":
    analyze_template()