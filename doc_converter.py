#!/usr/bin/env python3

import os
import tempfile
from typing import Optional, Tuple
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import streamlit as st

class DocConverter:
    """Handles conversion of legacy .doc files to .docx format for processing"""
    
    def __init__(self):
        self.temp_dir = tempfile.mkdtemp()
    
    def is_legacy_doc_file(self, file_path: str) -> bool:
        """Check if a file is a legacy .doc file that needs conversion"""
        if not file_path.endswith('.doc'):
            return False
        
        try:
            # Try to open with python-docx first
            Document(file_path)
            return False  # It's already in modern format
        except Exception as e:
            # If it fails with specific legacy format errors, it's a true .doc file
            error_str = str(e).lower()
            return any(keyword in error_str for keyword in [
                'not a word file', 'content type', 'zip', 'ole', 'compound'
            ])
    
    def convert_doc_to_docx(self, doc_file_path: str) -> Optional[str]:
        """Convert legacy .doc file to .docx format using API text extraction"""
        # Method 1: Try API-based text extraction (best quality)
        try:
            result = self._convert_using_api_extraction(doc_file_path)
            if result:
                st.success("🔄 Conversion réussie avec extraction API")
                return result
        except Exception as e:
            st.warning(f"API extraction failed: {str(e)}")
            pass  # Continue to fallback methods
            
        # Method 2: Try LibreOffice conversion (fallback)
        try:
            result = self._convert_using_libreoffice(doc_file_path)
            if result:
                st.success("🔄 Conversion réussie avec LibreOffice")
                return result
        except Exception as e:
            pass  # Continue to next method
            
        # Method 3: Try antiword if available (Linux/Unix tool)
        try:
            result = self._convert_using_antiword(doc_file_path)
            if result:
                st.success("🔄 Conversion réussie avec antiword")
                return result
        except Exception as e:
            pass  # Continue to next method
            
        # Method 4: Try python-docx2txt (may work for some .doc files)
        try:
            result = self._convert_using_docx2txt(doc_file_path)
            if result:
                st.info("🔄 Conversion basique réussie (qualité limitée)")
                return result
        except Exception as e:
            pass  # Continue to next method
            
        # Method 5: Basic text extraction as last resort
        try:
            result = self._convert_using_basic_extraction(doc_file_path)
            if result:
                st.warning("⚠️ Conversion basique - formatage minimal préservé")
                return result
        except Exception as e:
            pass
    
    def _convert_using_api_extraction(self, doc_file_path: str) -> str:
        """Convert .doc file using OpenAI API for clean text extraction"""
        import requests
        import base64
        
        # Get API key
        api_key = self._get_api_key()
        if not api_key:
            raise ValueError("Clé API OpenAI non trouvée")
        
        # Read and encode file
        with open(doc_file_path, 'rb') as f:
            file_content = f.read()
        
        # Use OpenAI API to extract text from the binary content
        payload = {
            "model": "gpt-4o",
            "messages": [
                {
                    "role": "system", 
                    "content": "Tu es un expert en extraction de texte de documents Word legacy. Extrais seulement le texte principal du document, en ignorant les numéros de page, en-têtes, pieds de page, et autres éléments de mise en forme. Conserve la structure des paragraphes et les sauts de ligne naturels. Réponds uniquement avec le texte extrait, sans commentaire."
                },
                {
                    "role": "user", 
                    "content": f"Voici le contenu binaire d'un fichier Word .doc legacy. Extrais le texte principal:\n\nContenu (premiers 8000 caractères): {str(file_content[:8000])}"
                }
            ],
            "temperature": 0,
            "max_tokens": 4000
        }
        
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        response = requests.post(
            "https://api.openai.com/v1/chat/completions",
            json=payload,
            headers=headers,
            timeout=60
        )
        
        if response.status_code != 200:
            raise ValueError(f"Erreur API: {response.status_code}")
        
        result = response.json()
        extracted_text = result.get("choices", [{}])[0].get("message", {}).get("content", "")
        
        if not extracted_text or len(extracted_text.strip()) < 10:
            raise ValueError("Texte extrait insuffisant")
        
        # Create .docx with extracted text
        return self._create_docx_from_text(extracted_text, doc_file_path)
    
    def _get_api_key(self) -> str:
        """Get OpenAI API key from various sources"""
        # Try Streamlit secrets first
        try:
            import streamlit as st
            if hasattr(st, 'secrets') and 'OPENAI_API_KEY' in st.secrets:
                key = st.secrets['OPENAI_API_KEY']
                if key and key.strip():
                    return key.strip()
        except Exception:
            pass
        
        # Try environment variable
        try:
            import os
            key = os.getenv('OPENAI_API_KEY')
            if key and key.strip():
                return key.strip()
        except Exception:
            pass
        
        # Try local file
        import os
        candidates = [
            os.path.join(os.getcwd(), 'cleAPI.txt'),
            os.path.abspath(os.path.join(os.path.dirname(__file__), 'cleAPI.txt')),
        ]
        for p in candidates:
            try:
                if os.path.exists(p):
                    with open(p, 'r', encoding='utf-8') as f:
                        key = f.read().strip()
                        if key:
                            return key
            except Exception:
                continue
        return ""
    
    def _convert_using_docx2txt(self, doc_file_path: str) -> str:
        """Convert using python-docx2txt library"""
        try:
            import docx2txt
        except ImportError:
            raise ImportError("python-docx2txt not available")
        
        # Extract text from .doc file
        text_content = docx2txt.process(doc_file_path)
        
        if not text_content or not text_content.strip():
            raise ValueError("Aucun texte extrait du fichier .doc")
        
        # Create a new .docx file with the extracted content
        new_doc = Document()
        
        # Split content into paragraphs and add them
        paragraphs = text_content.split('\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = new_doc.add_paragraph(para_text.strip())
                # Apply basic formatting
                for run in para.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
        
        # Save to temporary file
        temp_docx_path = os.path.join(self.temp_dir, f"converted_{os.path.basename(doc_file_path)}.docx")
        new_doc.save(temp_docx_path)
        
        return temp_docx_path
    
    def _convert_using_antiword(self, doc_file_path: str) -> str:
        """Convert using antiword (Linux/Unix tool for .doc files)"""
        import subprocess
        import shutil
        
        # Check if antiword is available
        if not shutil.which('antiword'):
            raise ValueError("antiword non trouvé sur le système")
        
        try:
            # Use antiword to extract text
            result = subprocess.run(
                ['antiword', doc_file_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            
            if result.returncode != 0:
                raise ValueError(f"Conversion antiword échouée: {result.stderr}")
            
            text_content = result.stdout.strip()
            if not text_content:
                raise ValueError("Aucun texte extrait avec antiword")
            
            return self._create_docx_from_text(text_content, doc_file_path)
            
        except subprocess.TimeoutExpired:
            raise ValueError("Timeout lors de la conversion antiword")
        except Exception as e:
            raise ValueError(f"Erreur antiword: {str(e)}")
    
    def _convert_using_basic_extraction(self, doc_file_path: str) -> str:
        """Basic text extraction as last resort"""
        try:
            with open(doc_file_path, 'rb') as f:
                content = f.read()
            
            # Try to extract readable text
            text_content = self._extract_text_from_binary(content)
            
            if not text_content or len(text_content.strip()) < 10:
                raise ValueError("Pas assez de texte extrait")
            
            return self._create_docx_from_text(text_content, doc_file_path)
            
        except Exception as e:
            raise ValueError(f"Extraction basique échouée: {str(e)}")
    
    def _convert_using_libreoffice(self, doc_file_path: str) -> str:
        """Convert using LibreOffice command line (explicit Word filter)"""
        import subprocess
        import shutil
        
        # Check if LibreOffice is available
        libreoffice_cmd = None
        for cmd in ['libreoffice', 'soffice']:
            if shutil.which(cmd):
                libreoffice_cmd = cmd
                break
        
        if not libreoffice_cmd:
            raise ValueError("LibreOffice non trouvé sur le système")
        
        # Convert using LibreOffice headless mode
        output_dir = tempfile.mkdtemp()
        cmd = [
            libreoffice_cmd,
            '--headless',
            '--convert-to', 'docx:"MS Word 2007 XML"',
            '--outdir', output_dir,
            doc_file_path
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        
        if result.returncode != 0:
            raise ValueError(f"Conversion LibreOffice échouée: {result.stderr or result.stdout}")
        
        # Find the converted file
        base_name = os.path.splitext(os.path.basename(doc_file_path))[0]
        converted_path = os.path.join(output_dir, f"{base_name}.docx")
        
        if not os.path.exists(converted_path):
            raise ValueError("Fichier converti non trouvé")
        
        # Move to our temp directory
        final_path = os.path.join(self.temp_dir, f"converted_{base_name}.docx")
        shutil.move(converted_path, final_path)
        shutil.rmtree(output_dir, ignore_errors=True)
        
        return final_path
    
    def _extract_text_from_binary(self, binary_content: bytes) -> str:
        """Extract readable text from binary .doc content (improved approach)"""
        # This is a more sophisticated text extraction approach
        text_chars = []
        
        for byte in binary_content:
            # Look for printable ASCII characters and extended ASCII for accents
            if 32 <= byte <= 126 or 128 <= byte <= 255:  # Extended ASCII range
                text_chars.append(chr(byte))
            elif byte in [10, 13]:  # Newline characters
                text_chars.append('\n')
            else:
                # Skip control characters but keep spacing
                if text_chars and text_chars[-1] not in [' ', '\n']:
                    text_chars.append(' ')
        
        # Clean up the extracted text
        raw_text = ''.join(text_chars)
        
        # Better text cleaning
        lines = []
        for line in raw_text.split('\n'):
            # Remove excessive spaces and clean up
            cleaned_line = ' '.join(line.split())
            
            # Filter out lines that are likely formatting artifacts
            if (len(cleaned_line) > 5 and  # Minimum length
                not cleaned_line.startswith(('>', '<', '{', '}', '\\', 'x')) and  # Skip formatting chars
                not all(c in '~=-_*+#@$%^&()[]{}|\\<>/?.,;:\'"`' for c in cleaned_line) and  # Skip symbol lines
                any(c.isalpha() for c in cleaned_line)):  # Must contain letters
                lines.append(cleaned_line)
        
        # Join with proper paragraph breaks
        return '\n\n'.join(lines)
    
    def _create_docx_from_text(self, text_content: str, original_path: str) -> str:
        """Create a .docx file from extracted text"""
        new_doc = Document()
        
        # Split content into paragraphs and add them
        paragraphs = text_content.split('\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = new_doc.add_paragraph(para_text.strip())
        
        # Save to temporary file
        base_name = os.path.splitext(os.path.basename(original_path))[0]
        temp_docx_path = os.path.join(self.temp_dir, f"converted_{base_name}.docx")
        new_doc.save(temp_docx_path)
        
        return temp_docx_path
    
    def get_document_info(self, file_path: str) -> Tuple[bool, str]:
        """Get information about document format and conversion needs"""
        if file_path.endswith('.docx'):
            return False, "Format moderne .docx"
        
        if file_path.endswith('.doc'):
            if self.is_legacy_doc_file(file_path):
                return True, "Format legacy .doc (Word 97-2003) - conversion requise"
            else:
                return False, "Format .doc moderne compatible"
        
        return False, "Format non Word"
    
    def cleanup(self):
        """Clean up temporary files"""
        try:
            import shutil
            shutil.rmtree(self.temp_dir, ignore_errors=True)
        except Exception:
            pass