#!/usr/bin/env python3

import os
import tempfile
from typing import Optional, Tuple
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import streamlit as st

class DocConverter:
    """Handles conversion of legacy .doc files to .docx format for processing"""
    
    def __init__(self):
        self.temp_dir = tempfile.mkdtemp()
    
    def is_legacy_doc_file(self, file_path: str) -> bool:
        """Check if a file is a legacy .doc file that needs conversion"""
        if not file_path.endswith('.doc'):
            return False
        
        try:
            # Try to open with python-docx first
            Document(file_path)
            return False  # It's already in modern format
        except Exception as e:
            # If it fails with specific legacy format errors, it's a true .doc file
            error_str = str(e).lower()
            return any(keyword in error_str for keyword in [
                'not a word file', 'content type', 'zip', 'ole', 'compound'
            ])
    
    def convert_doc_to_docx(self, doc_file_path: str) -> Optional[str]:
        """Convert legacy .doc file to .docx format - for now, create a placeholder"""
        
        # For problematic .doc files, create a clear placeholder document
        st.error(f"⚠️ Fichier .doc non supporté: {os.path.basename(doc_file_path)}")
        st.info("💡 Solution: Convertir manuellement ce fichier .doc en .docx avec Microsoft Word ou LibreOffice")
        
        # Create a placeholder document explaining the issue
        new_doc = Document()
        
        # Add a clear message explaining the situation
        title_para = new_doc.add_paragraph()
        title_run = title_para.add_run("⚠️ FICHIER NON CONVERTI")
        title_run.font.bold = True
        title_run.font.size = Pt(14)
        title_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # Red color
        
        new_doc.add_paragraph("")  # Empty line
        
        info_para = new_doc.add_paragraph()
        info_run = info_para.add_run(f"Le fichier '{os.path.basename(doc_file_path)}' est au format .doc legacy et n'a pas pu être converti automatiquement.")
        info_run.font.size = Pt(11)
        
        new_doc.add_paragraph("")  # Empty line
        
        solution_para = new_doc.add_paragraph()
        solution_run = solution_para.add_run("Solution recommandée :")
        solution_run.font.bold = True
        solution_run.font.size = Pt(11)
        
        steps = [
            "1. Ouvrir le fichier .doc avec Microsoft Word ou LibreOffice",
            "2. Faire 'Enregistrer sous' et choisir le format .docx",
            "3. Remplacer le fichier .doc par le nouveau .docx dans le dossier des clauses",
            "4. Recharger l'application"
        ]
        
        for step in steps:
            step_para = new_doc.add_paragraph()
            step_run = step_para.add_run(step)
            step_run.font.size = Pt(10)
        
        # Save placeholder document
        temp_docx_path = os.path.join(self.temp_dir, f"placeholder_{os.path.basename(doc_file_path)}.docx")
        new_doc.save(temp_docx_path)
        
        return temp_docx_path
    
    def _convert_using_textract(self, doc_file_path: str) -> str:
        """Convert using mammoth library for clean text extraction"""
        try:
            import mammoth
        except ImportError:
            raise ImportError("mammoth not available")
        
        try:
            # Extract text using mammoth
            with open(doc_file_path, "rb") as docx_file:
                result = mammoth.extract_raw_text(docx_file)
                text_content = result.value
        except Exception as e:
            # Fallback: try to convert .doc to .docx first
            temp_docx = self._try_doc_to_docx_conversion(doc_file_path)
            if temp_docx:
                with open(temp_docx, "rb") as docx_file:
                    result = mammoth.extract_raw_text(docx_file)
                    text_content = result.value
            else:
                raise e
        
        if not text_content or not text_content.strip():
            raise ValueError("Aucun texte extrait du fichier")
        
        # Clean the text - remove excessive whitespace and normalize
        lines = []
        for line in text_content.split('\n'):
            cleaned_line = line.strip()
            if cleaned_line:
                lines.append(cleaned_line)
        
        # Join lines with proper spacing
        clean_text = '\n\n'.join(lines)
        
        # Create a new .docx file with the clean extracted text
        new_doc = Document()
        
        # Add each paragraph to the document
        paragraphs = clean_text.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = new_doc.add_paragraph(para_text.strip())
                # Apply basic formatting
                for run in para.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
        
        # Save to temporary file
        temp_docx_path = os.path.join(self.temp_dir, f"mammoth_converted_{os.path.basename(doc_file_path)}.docx")
        new_doc.save(temp_docx_path)
        
        return temp_docx_path
    
    def _try_doc_to_docx_conversion(self, doc_file_path: str) -> Optional[str]:
        """Try to convert .doc to .docx for mammoth processing"""
        try:
            return self._convert_using_libreoffice(doc_file_path)
        except:
            try:
                return self._convert_using_antiword(doc_file_path)
            except:
                return None
    
    def _convert_using_docx2txt(self, doc_file_path: str) -> str:
        """Convert using python-docx2txt library - but only if it's really a .docx disguised as .doc"""
        try:
            import docx2txt
        except ImportError:
            raise ImportError("python-docx2txt not available")
        
        # Check if it's really a .docx file disguised as .doc
        try:
            # Try to read as zip (docx files are zip archives)
            import zipfile
            with zipfile.ZipFile(doc_file_path, 'r') as zip_ref:
                if 'word/document.xml' in zip_ref.namelist():
                    # It's actually a .docx file, so docx2txt will work
                    text_content = docx2txt.process(doc_file_path)
                else:
                    raise ValueError("Not a .docx file in disguise")
        except (zipfile.BadZipFile, ValueError):
            # It's a real .doc file, docx2txt won't work
            raise ValueError("Cannot process true .doc files with docx2txt")
        
        if not text_content or not text_content.strip():
            raise ValueError("Aucun texte extrait du fichier .doc")
        
        # Create a new .docx file with the extracted content
        new_doc = Document()
        
        # Split content into paragraphs and add them
        paragraphs = text_content.split('\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = new_doc.add_paragraph(para_text.strip())
                # Apply basic formatting
                for run in para.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
        
        # Save to temporary file
        temp_docx_path = os.path.join(self.temp_dir, f"converted_{os.path.basename(doc_file_path)}.docx")
        new_doc.save(temp_docx_path)
        
        return temp_docx_path
    
    def _convert_using_antiword(self, doc_file_path: str) -> str:
        """Convert using antiword (Linux/Unix tool for .doc files)"""
        import subprocess
        import shutil
        
        # Check if antiword is available
        if not shutil.which('antiword'):
            raise ValueError("antiword non trouvé sur le système")
        
        try:
            # Use antiword to extract text
            result = subprocess.run(
                ['antiword', doc_file_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            
            if result.returncode != 0:
                raise ValueError(f"Conversion antiword échouée: {result.stderr}")
            
            text_content = result.stdout.strip()
            if not text_content:
                raise ValueError("Aucun texte extrait avec antiword")
            
            return self._create_docx_from_text(text_content, doc_file_path)
            
        except subprocess.TimeoutExpired:
            raise ValueError("Timeout lors de la conversion antiword")
        except Exception as e:
            raise ValueError(f"Erreur antiword: {str(e)}")
    
    def _convert_using_basic_extraction(self, doc_file_path: str) -> str:
        """Basic text extraction as last resort"""
        try:
            with open(doc_file_path, 'rb') as f:
                content = f.read()
            
            # Try to extract readable text
            text_content = self._extract_text_from_binary(content)
            
            if not text_content or len(text_content.strip()) < 10:
                raise ValueError("Pas assez de texte extrait")
            
            return self._create_docx_from_text(text_content, doc_file_path)
            
        except Exception as e:
            raise ValueError(f"Extraction basique échouée: {str(e)}")
    
    def _convert_using_libreoffice(self, doc_file_path: str) -> str:
        """Convert using LibreOffice command line (explicit Word filter)"""
        import subprocess
        import shutil
        
        # Check if LibreOffice is available
        libreoffice_cmd = None
        for cmd in ['libreoffice', 'soffice']:
            if shutil.which(cmd):
                libreoffice_cmd = cmd
                break
        
        if not libreoffice_cmd:
            raise ValueError("LibreOffice non trouvé sur le système")
        
        # Convert using LibreOffice headless mode
        output_dir = tempfile.mkdtemp()
        cmd = [
            libreoffice_cmd,
            '--headless',
            '--convert-to', 'docx:"MS Word 2007 XML"',
            '--outdir', output_dir,
            doc_file_path
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        
        if result.returncode != 0:
            raise ValueError(f"Conversion LibreOffice échouée: {result.stderr or result.stdout}")
        
        # Find the converted file
        base_name = os.path.splitext(os.path.basename(doc_file_path))[0]
        converted_path = os.path.join(output_dir, f"{base_name}.docx")
        
        if not os.path.exists(converted_path):
            raise ValueError("Fichier converti non trouvé")
        
        # Move to our temp directory
        final_path = os.path.join(self.temp_dir, f"converted_{base_name}.docx")
        shutil.move(converted_path, final_path)
        shutil.rmtree(output_dir, ignore_errors=True)
        
        return final_path
    
    def _extract_text_from_binary(self, binary_content: bytes) -> str:
        """Extract readable text from binary .doc content (improved approach)"""
        # This is a more sophisticated text extraction approach
        text_chars = []
        
        for byte in binary_content:
            # Look for printable ASCII characters and extended ASCII for accents
            if 32 <= byte <= 126 or 128 <= byte <= 255:  # Extended ASCII range
                text_chars.append(chr(byte))
            elif byte in [10, 13]:  # Newline characters
                text_chars.append('\n')
            else:
                # Skip control characters but keep spacing
                if text_chars and text_chars[-1] not in [' ', '\n']:
                    text_chars.append(' ')
        
        # Clean up the extracted text
        raw_text = ''.join(text_chars)
        
        # Better text cleaning
        lines = []
        for line in raw_text.split('\n'):
            # Remove excessive spaces and clean up
            cleaned_line = ' '.join(line.split())
            
            # Filter out lines that are likely formatting artifacts
            if (len(cleaned_line) > 5 and  # Minimum length
                not cleaned_line.startswith(('>', '<', '{', '}', '\\', 'x')) and  # Skip formatting chars
                not all(c in '~=-_*+#@$%^&()[]{}|\\<>/?.,;:\'"`' for c in cleaned_line) and  # Skip symbol lines
                any(c.isalpha() for c in cleaned_line)):  # Must contain letters
                lines.append(cleaned_line)
        
        # Join with proper paragraph breaks
        return '\n\n'.join(lines)
    
    def _create_docx_from_text(self, text_content: str, original_path: str) -> str:
        """Create a .docx file from extracted text"""
        new_doc = Document()
        
        # Split content into paragraphs and add them
        paragraphs = text_content.split('\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = new_doc.add_paragraph(para_text.strip())
        
        # Save to temporary file
        base_name = os.path.splitext(os.path.basename(original_path))[0]
        temp_docx_path = os.path.join(self.temp_dir, f"converted_{base_name}.docx")
        new_doc.save(temp_docx_path)
        
        return temp_docx_path
    
    def get_document_info(self, file_path: str) -> Tuple[bool, str]:
        """Get information about document format and conversion needs"""
        if file_path.endswith('.docx'):
            return False, "Format moderne .docx"
        
        if file_path.endswith('.doc'):
            if self.is_legacy_doc_file(file_path):
                return True, "Format legacy .doc (Word 97-2003) - conversion requise"
            else:
                return False, "Format .doc moderne compatible"
        
        return False, "Format non Word"
    
    def cleanup(self):
        """Clean up temporary files"""
        try:
            import shutil
            shutil.rmtree(self.temp_dir, ignore_errors=True)
        except Exception:
            pass